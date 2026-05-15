# =============================================================================
# update_manuals.py
# 사랑과평안의교회 - 부서 매뉴얼 자동 업데이트 스크립트
#
# Requirements (requirements.txt 내용):
#   python-docx==1.1.2
#   requests==2.32.3
#   google-auth==2.29.0
#   google-auth-httplib2==0.2.0
#   google-api-python-client==2.127.0
# =============================================================================

from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# 설정 상수
# =============================================================================

# DB 폴더 경로
DB_FOLDER = Path(r"E:\사평교 메뉴얼 제정\DB")

# 교회 이름
CHURCH_NAME = "사랑과평안의교회"

# Google Sheets API 키 또는 서비스 계정 JSON 경로 (환경변수로 관리)
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")
GOOGLE_SERVICE_ACCOUNT_JSON = os.environ.get(
    "GOOGLE_SERVICE_ACCOUNT_JSON", ""
)

# 부서명 → docx 파일명 매핑
DEPARTMENT_FILES: dict[str, str] = {
    "예배부": "WOR_예배부매뉴얼_201-203_사랑과평안의교회.docx",
    "재정부": "FIN_재정부매뉴얼_201-203_사랑과평안의교회.docx",
    "새가족부": "NEW_새가족부매뉴얼_201-203_사랑과평안의교회.docx",
    "교육부": "EDU_교육부매뉴얼_201-204_사랑과평안의교회.docx",
    "친교봉사부": "FEL_친교봉사부매뉴얼_201-203_사랑과평안의교회.docx",
    "시설관리부": "FAC_시설관리부매뉴얼_201-203_사랑과평안의교회.docx",
    "환경미화부": "ENV_환경미화부매뉴얼_201-203_사랑과평안의교회.docx",
    "경조사부": "CEL_경조사부매뉴얼_201-203_사랑과평안의교회.docx",
    "문화사역팀": "CUL_문화사역팀매뉴얼_201-203_사랑과평안의교회.docx",
    "문서사역팀": "DOC_문서사역팀매뉴얼_201-203_사랑과평안의교회.docx",
    "성가대": "CHO_성가대매뉴얼_201-203_사랑과평안의교회.docx",
    "찬양팀": "PRS_찬양팀매뉴얼_201-203_사랑과평안의교회.docx",
}

# 업무현황 섹션을 찾을 때 사용하는 후보 키워드 목록
WORK_SECTION_KEYWORDS = [
    "업무현황",
    "업무 현황",
    "업무 범위",
    "업무범위",
    "제 2 장",
    "제2장",
]


# =============================================================================
# Google Sheets 읽기
# =============================================================================

def read_google_sheet(sheet_id: str, sheet_name: str) -> list[dict[str, Any]]:
    """
    Google Sheets에서 데이터를 읽어 딕셔너리 리스트로 반환한다.

    서비스 계정 JSON(GOOGLE_SERVICE_ACCOUNT_JSON 환경변수) 또는
    API 키(GOOGLE_API_KEY 환경변수) 방식을 지원한다.
    첫 번째 행을 헤더로 사용한다.

    Args:
        sheet_id: Google Sheets 파일 ID (URL에서 /d/{id}/ 부분)
        sheet_name: 읽을 시트 이름 (예: "업무현황")

    Returns:
        각 행을 딕셔너리로 변환한 리스트
        예: [{"부서": "예배부", "업무명": "주일예배 진행", ...}, ...]

    Raises:
        ValueError: API 키 또는 서비스 계정 정보가 없을 때
        requests.HTTPError: API 요청 실패 시
    """
    # 서비스 계정 방식 시도
    if GOOGLE_SERVICE_ACCOUNT_JSON and Path(GOOGLE_SERVICE_ACCOUNT_JSON).exists():
        return _read_sheet_with_service_account(sheet_id, sheet_name)

    # API 키 방식 시도
    if GOOGLE_API_KEY:
        return _read_sheet_with_api_key(sheet_id, sheet_name)

    raise ValueError(
        "Google Sheets 접근 정보가 없습니다.\n"
        "환경변수 GOOGLE_API_KEY 또는 GOOGLE_SERVICE_ACCOUNT_JSON을 설정하세요."
    )


def _read_sheet_with_api_key(sheet_id: str, sheet_name: str) -> list[dict[str, Any]]:
    """API 키를 사용하여 공개 Google Sheet 데이터를 읽는다."""
    # URL 인코딩: 시트 이름에 한글이 있을 수 있음
    encoded_range = requests.utils.quote(f"{sheet_name}!A1:Z1000")
    url = (
        f"https://sheets.googleapis.com/v4/spreadsheets/{sheet_id}"
        f"/values/{encoded_range}?key={GOOGLE_API_KEY}"
    )

    response = requests.get(url, timeout=30)
    response.raise_for_status()

    data = response.json()
    rows: list[list[str]] = data.get("values", [])

    if not rows:
        print(f"  [경고] 시트 '{sheet_name}'에 데이터가 없습니다.")
        return []

    return _rows_to_dicts(rows)


def _read_sheet_with_service_account(
    sheet_id: str, sheet_name: str
) -> list[dict[str, Any]]:
    """서비스 계정 JSON을 사용하여 Google Sheet 데이터를 읽는다."""
    try:
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
    except ImportError as exc:
        raise ImportError(
            "google-auth, google-api-python-client 패키지가 필요합니다.\n"
            "pip install google-auth google-api-python-client"
        ) from exc

    scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
    credentials = service_account.Credentials.from_service_account_file(
        GOOGLE_SERVICE_ACCOUNT_JSON, scopes=scopes
    )
    service = build("sheets", "v4", credentials=credentials)

    range_name = f"{sheet_name}!A1:Z1000"
    result = (
        service.spreadsheets()
        .values()
        .get(spreadsheetId=sheet_id, range=range_name)
        .execute()
    )
    rows: list[list[str]] = result.get("values", [])

    if not rows:
        print(f"  [경고] 시트 '{sheet_name}'에 데이터가 없습니다.")
        return []

    return _rows_to_dicts(rows)


def _rows_to_dicts(rows: list[list[str]]) -> list[dict[str, Any]]:
    """
    2차원 행 데이터를 헤더를 키로 사용하는 딕셔너리 리스트로 변환한다.

    첫 번째 행이 헤더이며, 열 수가 헤더보다 짧은 행은 빈 문자열로 패딩한다.
    """
    if len(rows) < 2:
        return []

    headers = [str(h).strip() for h in rows[0]]
    result: list[dict[str, Any]] = []

    for row in rows[1:]:
        # 행이 헤더보다 짧으면 빈 문자열로 채움
        padded = row + [""] * (len(headers) - len(row))
        record = {headers[i]: padded[i] for i in range(len(headers))}

        # 완전히 빈 행 건너뜀
        if any(v.strip() for v in record.values()):
            result.append(record)

    return result


# =============================================================================
# docx 섹션 탐색
# =============================================================================

def find_section(doc: Document, search_text: str) -> int | None:
    """
    문서에서 search_text를 포함하는 단락의 인덱스를 반환한다.

    대소문자 및 공백을 정규화하여 검색한다.

    Args:
        doc: python-docx Document 객체
        search_text: 검색할 텍스트 (부분 일치)

    Returns:
        첫 번째로 일치하는 단락 인덱스, 없으면 None
    """
    normalized_search = search_text.replace(" ", "").lower()

    for idx, para in enumerate(doc.paragraphs):
        normalized_para = para.text.replace(" ", "").lower()
        if normalized_search in normalized_para:
            return idx

    return None


def find_work_section(doc: Document) -> int | None:
    """
    업무현황/업무 범위 섹션의 단락 인덱스를 반환한다.

    WORK_SECTION_KEYWORDS 목록을 순서대로 시도한다.

    Args:
        doc: python-docx Document 객체

    Returns:
        섹션 단락 인덱스, 없으면 None
    """
    for keyword in WORK_SECTION_KEYWORDS:
        idx = find_section(doc, keyword)
        if idx is not None:
            return idx

    return None


# =============================================================================
# docx 내용 업데이트
# =============================================================================

def _make_bold_paragraph(doc: Document, text: str, level: int = 0) -> None:
    """
    문서에 굵은 텍스트 단락을 추가한다.

    Args:
        doc: 대상 Document 객체 (직접 append 방식 사용 불가 → _insert_after 참고)
        text: 단락 텍스트
        level: 들여쓰기 수준 (0 = 없음)
    """
    para = doc.add_paragraph()
    if level > 0:
        para.paragraph_format.left_indent = Pt(18 * level)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return para


def _make_bullet_paragraph(doc: Document, text: str) -> None:
    """
    문서에 불릿 항목 단락을 추가한다.

    Args:
        doc: 대상 Document 객체
        text: 불릿 항목 텍스트
    """
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    return para


def _insert_paragraph_after(reference_para, new_para) -> None:
    """
    reference_para 뒤에 new_para를 삽입한다.

    python-docx는 특정 위치 삽입을 직접 지원하지 않으므로
    XML 조작을 사용한다.
    """
    ref_element = reference_para._element
    new_element = new_para._element
    ref_element.addnext(new_element)


def update_work_section(
    doc: Document,
    work_data: list[dict[str, Any]],
    preview: bool = False,
) -> list[str]:
    """
    문서의 업무현황 섹션을 work_data로 업데이트한다.

    기존 섹션 바로 뒤에 새 내용을 삽입하거나, 섹션이 없으면
    문서 끝에 섹션 제목과 함께 추가한다.

    업무 항목 형식:
        [굵은 글씨] 업무명
          • 세부 내용 1
          • 세부 내용 2
          ...

    Args:
        doc: 수정할 Document 객체
        work_data: 업무 데이터 딕셔너리 리스트
                   각 항목에 '업무명', '세부내용' 키가 있어야 함
        preview: True이면 변경 내용만 출력하고 문서를 수정하지 않음

    Returns:
        변경 내용 요약 문자열 리스트 (preview 출력 또는 로그용)
    """
    changes: list[str] = []

    for item in work_data:
        work_title = str(item.get("업무명", "")).strip()
        detail_raw = str(item.get("세부내용", "")).strip()

        if not work_title:
            continue

        # 세부내용: 줄바꿈 또는 '/'로 분리
        details = [
            d.strip()
            for d in detail_raw.replace("/", "\n").splitlines()
            if d.strip()
        ]

        changes.append(f"  + 업무명: {work_title}")
        for d in details:
            changes.append(f"      • {d}")

    if preview or not changes:
        return changes

    # 업무현황 섹션 위치 탐색
    section_idx = find_work_section(doc)

    if section_idx is not None:
        # 섹션 제목 다음 단락부터 다음 챕터 전까지 범위를 확인하여
        # 마지막 단락 참조를 구한 뒤 그 뒤에 삽입
        anchor_para = doc.paragraphs[section_idx]
    else:
        # 섹션이 없으면 새로 생성하여 문서 끝에 추가
        doc.add_paragraph()  # 빈 줄
        anchor_para = _make_bold_paragraph(doc, "업무 현황")
        changes.insert(0, "  [섹션 새로 생성] 업무 현황")

    # 각 업무 항목을 anchor_para 이후에 순서대로 삽입
    # (역순으로 삽입하면 순서가 유지됨)
    insert_paras = []
    for item in work_data:
        work_title = str(item.get("업무명", "")).strip()
        detail_raw = str(item.get("세부내용", "")).strip()

        if not work_title:
            continue

        details = [
            d.strip()
            for d in detail_raw.replace("/", "\n").splitlines()
            if d.strip()
        ]

        # 임시로 문서 끝에 단락 생성 후 XML 이동 방식 사용
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(work_title)
        title_run.bold = True
        title_run.font.size = Pt(11)

        bullet_paras = []
        for detail in details:
            try:
                bp = doc.add_paragraph(style="List Bullet")
            except KeyError:
                # List Bullet 스타일이 없는 문서 대응
                bp = doc.add_paragraph()
                bp.paragraph_format.left_indent = Pt(18)
            run = bp.add_run(f"• {detail}")
            run.font.size = Pt(10.5)
            bullet_paras.append(bp)

        insert_paras.append((title_para, bullet_paras))

    # 역순으로 anchor_para 뒤에 삽입하여 올바른 순서 유지
    for title_para, bullet_paras in reversed(insert_paras):
        for bp in reversed(bullet_paras):
            _insert_paragraph_after(anchor_para, bp)
        _insert_paragraph_after(anchor_para, title_para)

    return changes


# =============================================================================
# 수정 이력 관리
# =============================================================================

def _increment_version(doc: Document) -> None:
    """
    문서에서 'Rev. X.Y' 패턴을 찾아 마이너 버전을 1 증가시킨다.

    예: Rev. 1.0 → Rev. 1.1, Rev. 1.9 → Rev. 1.10
    패턴이 없으면 첫 번째 단락 뒤에 'Rev. 1.0'을 추가한다.
    """
    import re
    rev_pattern = re.compile(r"Rev\.\s*(\d+)\.(\d+)", re.IGNORECASE)

    for para in doc.paragraphs:
        for run in para.runs:
            match = rev_pattern.search(run.text)
            if match:
                major = match.group(1)
                minor = int(match.group(2)) + 1
                run.text = rev_pattern.sub(f"Rev. {major}.{minor}", run.text)
                return

    # 버전 표시가 없으면 문서 두 번째 위치에 추가
    ver_para = doc.add_paragraph()
    ver_run = ver_para.add_run("Rev. 1.0")
    ver_run.font.size = Pt(9)
    ver_run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)
    if len(doc.element.body) > 1:
        doc.element.body.insert(1, ver_para._element)


def add_revision_history(
    doc: Document,
    changes: list[str],
    submitter: str,
) -> None:
    """
    문서 맨 앞에 수정 이력 테이블을 추가하거나 기존 테이블에 행을 추가한다.

    테이블 컬럼: 번호 | 수정일 | 수정 조항 | 내용 | 수정자

    이미 수정 이력 테이블이 있으면(첫 열 헤더가 "번호") 새 행만 추가한다.
    없으면 "수정 이력" 제목과 테이블을 문서 앞에 새로 삽입한다.

    마지막으로 _increment_version()을 호출해 버전 번호를 올린다.

    Args:
        doc: 수정할 Document 객체
        changes: update_work_section이 반환한 변경 요약 리스트
        submitter: 제출자 이름 (부서장 또는 차장)
    """
    from docx.oxml import OxmlElement

    today = datetime.now().strftime("%Y-%m-%d")

    # 변경된 업무명 추출 → 내용 요약
    work_titles = [
        c.replace("  + 업무명: ", "").strip()
        for c in changes
        if "업무명:" in c
    ]
    content_summary = ", ".join(work_titles) if work_titles else "업무현황 업데이트"
    article = "업무현황 섹션"

    # ── 기존 수정 이력 테이블 탐색 ────────────────────────────────────────
    existing_table = None
    if doc.tables:
        first_tbl = doc.tables[0]
        if first_tbl.rows and first_tbl.rows[0].cells[0].text.strip() == "번호":
            existing_table = first_tbl

    if existing_table is not None:
        # 기존 테이블에 행 추가
        row_num = len(existing_table.rows)  # 헤더 포함 → 데이터 행 번호
        new_row = existing_table.add_row()
        for i, val in enumerate([str(row_num), today, article, content_summary, submitter]):
            new_row.cells[i].text = val
            runs = new_row.cells[i].paragraphs[0].runs
            if runs:
                runs[0].font.size = Pt(9)
    else:
        # ── 새 이력 테이블 생성 ────────────────────────────────────────────
        headers = ["번호", "수정일", "수정 조항", "내용", "수정자"]
        table = doc.add_table(rows=2, cols=len(headers))
        table.style = "Table Grid"

        # 헤더 행: 네이비 배경 + 흰 글씨
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1E3A8A")
            tcPr.append(shd)

        # 첫 번째 데이터 행
        for i, val in enumerate(["1", today, article, content_summary, submitter]):
            cell = table.rows[1].cells[i]
            cell.text = val
            runs = cell.paragraphs[0].runs
            if runs:
                runs[0].font.size = Pt(9)

        # 테이블과 제목을 문서 맨 앞에 이동 (XML 삽입)
        body = doc.element.body
        # 제목 단락
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("수정 이력")
        title_run.bold = True
        title_run.font.size = Pt(10)
        body.insert(0, table._tbl)
        body.insert(0, title_para._element)

    # 버전 번호 갱신
    _increment_version(doc)


# =============================================================================
# 부서별 처리
# =============================================================================

def _get_file_path(dept_name: str) -> Path:
    """
    부서명으로 docx 파일 경로를 반환한다.

    Raises:
        KeyError: 부서명이 DEPARTMENT_FILES에 없을 때
        FileNotFoundError: 파일이 DB_FOLDER에 없을 때
    """
    if dept_name not in DEPARTMENT_FILES:
        raise KeyError(
            f"알 수 없는 부서명: '{dept_name}'\n"
            f"지원 부서: {', '.join(DEPARTMENT_FILES.keys())}"
        )

    file_path = DB_FOLDER / DEPARTMENT_FILES[dept_name]

    if not file_path.exists():
        # DB_FOLDER에서 유사 파일 탐색하여 힌트 제공
        candidates = list(DB_FOLDER.glob("*.docx"))
        hint = f"\n  DB 폴더 내 파일: {[f.name for f in candidates]}" if candidates else ""
        raise FileNotFoundError(
            f"파일을 찾을 수 없습니다: {file_path}{hint}"
        )

    return file_path


def _backup_file(file_path: Path) -> Path:
    """
    원본 파일을 타임스탬프 백업 파일로 복사한다.

    형식: {stem}_backup_{YYYYMMDD}.docx
    같은 날짜에 여러 번 실행하면 기존 백업을 덮어쓴다.

    Returns:
        백업 파일 경로
    """
    timestamp = datetime.now().strftime("%Y%m%d")
    backup_name = f"{file_path.stem}_backup_{timestamp}{file_path.suffix}"
    backup_path = file_path.parent / backup_name
    shutil.copy2(file_path, backup_path)
    return backup_path


def process_department(
    dept_name: str,
    works_data: list[dict[str, Any]],
    preview: bool = False,
) -> dict[str, Any]:
    """
    단일 부서의 매뉴얼 docx 파일을 업데이트한다.

    1. 파일 경로 확인
    2. 백업 생성 (preview 모드가 아닐 때)
    3. 문서 열기
    4. 업무현황 섹션 업데이트
    5. 저장

    Args:
        dept_name: 부서명 (DEPARTMENT_FILES 키)
        works_data: 해당 부서의 업무 데이터 리스트
        preview: True이면 실제 저장 없이 변경 내용만 출력

    Returns:
        처리 결과 딕셔너리:
          {
            "dept": str,
            "count": int,       # 처리된 업무 항목 수
            "backup": str|None, # 백업 파일 경로
            "changes": list[str],
            "error": str|None,
          }
    """
    result: dict[str, Any] = {
        "dept": dept_name,
        "count": 0,
        "backup": None,
        "changes": [],
        "error": None,
    }

    try:
        file_path = _get_file_path(dept_name)
    except (KeyError, FileNotFoundError) as exc:
        result["error"] = str(exc)
        return result

    try:
        doc = Document(file_path)
    except Exception as exc:
        result["error"] = f"문서 열기 실패: {exc}"
        return result

    try:
        changes = update_work_section(doc, works_data, preview=preview)
        result["changes"] = changes
        result["count"] = sum(1 for item in works_data if item.get("업무명", "").strip())
    except Exception as exc:
        result["error"] = f"섹션 업데이트 실패: {exc}"
        return result

    # 수정 이력 테이블 추가 (preview 모드에서는 건너뜀)
    if not preview and changes:
        submitter = str(works_data[0].get("작성자", "")).strip() if works_data else ""
        if not submitter:
            submitter = "미기재"
        try:
            add_revision_history(doc, changes, submitter)
        except Exception as exc:
            # 이력 추가 실패는 치명적 오류가 아니므로 경고만 기록
            result["changes"].append(f"  [경고] 수정 이력 추가 실패: {exc}")

    if not preview:
        try:
            backup_path = _backup_file(file_path)
            result["backup"] = str(backup_path)
            doc.save(file_path)
        except Exception as exc:
            result["error"] = f"저장 실패: {exc}"
            return result

    return result


# =============================================================================
# 데이터 그룹핑
# =============================================================================

def group_by_department(
    rows: list[dict[str, Any]],
    dept_column: str = "부서",
) -> dict[str, list[dict[str, Any]]]:
    """
    행 리스트를 부서별로 그룹핑한다.

    Args:
        rows: Google Sheets에서 읽은 딕셔너리 리스트
        dept_column: 부서명이 저장된 열 이름 (기본값: "부서")

    Returns:
        {"예배부": [...], "재정부": [...], ...} 형태의 딕셔너리
    """
    grouped: dict[str, list[dict[str, Any]]] = {}
    for row in rows:
        dept = str(row.get(dept_column, "")).strip()
        if not dept:
            continue
        grouped.setdefault(dept, []).append(row)
    return grouped


# =============================================================================
# 메인 함수
# =============================================================================

def main(
    sheet_id: str,
    sheet_name: str = "업무현황",
    target_dept: str | None = None,
    preview: bool = False,
) -> None:
    """
    Google Sheets에서 제출된 데이터를 읽어 부서별 매뉴얼 docx를 업데이트한다.

    Args:
        sheet_id: Google Sheets 파일 ID
        sheet_name: 읽을 시트 이름
        target_dept: 특정 부서만 처리할 경우 부서명. None이면 전체 처리.
        preview: True이면 실제 파일을 수정하지 않고 변경 내용만 출력
    """
    print("=" * 60)
    print(f"  사랑과평안의교회 매뉴얼 업데이트")
    print(f"  실행 시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if preview:
        print("  [미리보기 모드] 파일이 수정되지 않습니다.")
    print("=" * 60)

    # Google Sheets 데이터 읽기
    print(f"\n[1/3] Google Sheets 데이터 읽기: 시트 '{sheet_name}' ...")
    try:
        rows = read_google_sheet(sheet_id, sheet_name)
    except Exception as exc:
        print(f"  [오류] Google Sheets 읽기 실패: {exc}")
        sys.exit(1)

    print(f"  → {len(rows)}개 행 읽음")

    if not rows:
        print("  처리할 데이터가 없습니다.")
        return

    # 부서별 그룹핑
    print("\n[2/3] 부서별 데이터 그룹핑 ...")
    grouped = group_by_department(rows)

    if target_dept:
        # 특정 부서만 필터링
        if target_dept not in grouped:
            print(f"  [경고] '{target_dept}' 부서 데이터가 없습니다.")
            print(f"  발견된 부서: {', '.join(grouped.keys()) or '없음'}")
            return
        grouped = {target_dept: grouped[target_dept]}

    print(f"  → 처리 대상 부서: {', '.join(grouped.keys())}")

    # 부서별 처리
    print("\n[3/3] 매뉴얼 파일 업데이트 ...")
    summary: list[str] = []

    for dept_name, dept_rows in grouped.items():
        print(f"\n  처리 중: {dept_name} ({len(dept_rows)}개 항목)")
        result = process_department(dept_name, dept_rows, preview=preview)

        if result["error"]:
            line = f"  [오류] {dept_name}: {result['error']}"
            print(line)
            summary.append(line)
            continue

        # 변경 내용 출력 (preview 또는 verbose 용)
        if result["changes"]:
            for change_line in result["changes"]:
                print(f"    {change_line}")

        if preview:
            line = f"  {dept_name}: {result['count']}개 업무 항목 미리보기 완료"
        else:
            backup_info = (
                f" (백업: {Path(result['backup']).name})"
                if result["backup"]
                else ""
            )
            line = f"  {dept_name}: {result['count']}개 업무 항목 업데이트 완료{backup_info}"

        print(line)
        summary.append(line)

    # 최종 요약
    print("\n" + "=" * 60)
    print("  [완료 요약]")
    for line in summary:
        print(line)
    print("=" * 60)


# =============================================================================
# CLI 인터페이스
# =============================================================================

def _build_arg_parser() -> argparse.ArgumentParser:
    """CLI 인수 파서를 생성하고 반환한다."""
    parser = argparse.ArgumentParser(
        description="사랑과평안의교회 부서 매뉴얼 자동 업데이트",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
사용 예시:
  # 모든 부서 업데이트
  python update_manuals.py --all --sheet-id YOUR_SHEET_ID

  # 특정 부서만 업데이트
  python update_manuals.py --dept 예배부 --sheet-id YOUR_SHEET_ID

  # 변경 내용 미리보기 (파일 수정 없음)
  python update_manuals.py --preview --sheet-id YOUR_SHEET_ID

  # 시트 이름 지정
  python update_manuals.py --all --sheet-id YOUR_SHEET_ID --sheet-name 업무현황2025

환경변수:
  GOOGLE_API_KEY               공개 시트용 API 키
  GOOGLE_SERVICE_ACCOUNT_JSON  서비스 계정 JSON 파일 경로 (비공개 시트)
        """,
    )

    mode_group = parser.add_mutually_exclusive_group(required=True)
    mode_group.add_argument(
        "--all",
        action="store_true",
        help="모든 부서 매뉴얼을 업데이트한다",
    )
    mode_group.add_argument(
        "--dept",
        metavar="부서명",
        help="특정 부서만 업데이트한다 (예: 예배부)",
    )
    mode_group.add_argument(
        "--preview",
        action="store_true",
        help="변경 내용을 미리보기만 한다 (파일 수정 없음)",
    )

    parser.add_argument(
        "--sheet-id",
        required=True,
        metavar="SHEET_ID",
        help="Google Sheets 파일 ID",
    )
    parser.add_argument(
        "--sheet-name",
        default="업무현황",
        metavar="SHEET_NAME",
        help="읽을 시트 이름 (기본값: 업무현황)",
    )

    return parser


if __name__ == "__main__":
    parser = _build_arg_parser()
    args = parser.parse_args()

    # 모드 결정
    is_preview = args.preview
    target_dept: str | None = None

    if args.dept:
        target_dept = args.dept
    elif args.preview:
        # --preview는 전체 부서에 대해 미리보기
        target_dept = None

    main(
        sheet_id=args.sheet_id,
        sheet_name=args.sheet_name,
        target_dept=target_dept,
        preview=is_preview,
    )
