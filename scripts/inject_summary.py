#!/usr/bin/env python3
"""복원된 원본 매뉴얼에 '제출 업무 총괄' 요약표를 병합 주입한다.

- base   : 작업트리의 복원 원본(.docx). 원본 상세는 그대로 둔다.
- source : snapshot/skeleton-2026-07-01 브랜치의 골격본에서 제출 업무총괄/예산 표를 추출.
- 주입   : 대표 매뉴얼(첫 '제1장') 앞에 네이비 헤더 + 요약표(업무명/주기/예산) 삽입.

사용: .venv/bin/python scripts/inject_summary.py
"""
import subprocess, io, glob, re, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SKELETON = "snapshot/skeleton-2026-07-01"
NAVY = "1A3A5C"; BORDER = "C0D4EC"
DEPTS = [("교육부","EDU"),("문화사역팀","CUL"),("새가족부","NEW"),("성가대","CHO"),
         ("시설관리부","FAC"),("예배부","WOR"),("재정부","FIN"),("찬양팀","PRS"),("환경미화부","ENV")]

def dept_file(d, code):
    for f in glob.glob(f"DB/{d}/{code}_*.docx"):
        if "_backup_orig" not in f and "_pilot" not in f:
            return f
    raise FileNotFoundError(d)

def skeleton_doc(path):
    b = subprocess.run(["git","show",f"{SKELETON}:{path}"], capture_output=True).stdout
    return Document(io.BytesIO(b))

def extract_summary(doc):
    """골격본에서 (업무명,주기,예산) 행과 항목별 예산 문자열을 추출."""
    rows, budget_items = [], []
    for t in doc.tables:
        cells = [[c.text.strip().replace("\n"," ") for c in r.cells] for r in t.rows]
        if not cells: continue
        hdr = " ".join(cells[0])
        if "업무명" in hdr and "주기" in hdr:
            for r in cells[1:]:
                if len(r) >= 3 and r[0]:
                    rows.append((r[0], r[1], r[2]))
        elif "항목" in hdr and ("금액" in hdr or "예산" in hdr):
            for r in cells[1:]:
                if len(r) >= 3 and r[0]:
                    budget_items.append(f"{r[0]} {r[1]} {r[2]}")
    budget_str = " · ".join(budget_items)
    # 예산 셀이 모호('참조'/공백/—)하고 단일 업무면 항목별 예산으로 보강
    if budget_str and len(rows) == 1:
        name, cyc, bud = rows[0]
        if (not bud) or bud in ("—","-") or "참조" in bud:
            rows[0] = (name, cyc, budget_str)
    return rows

def set_cell(cell, text, bold=False, bg=None, white=False):
    cell.text = ""
    r = cell.paragraphs[0].add_run(text)
    r.font.size = Pt(10); r.font.name = "맑은 고딕"; r.bold = bold
    if white: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    if bg:
        sh = OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),bg)
        cell._tc.get_or_add_tcPr().append(sh)

def add_borders(tbl):
    b = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"4")
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),BORDER); b.append(e)
    tbl._tbl.tblPr.append(b)

def count_articles(doc):
    return sum(1 for p in doc.paragraphs if re.match(r"^제\s*\d+\s*조", p.text.strip()))

def inject(dept, code):
    path = dept_file(dept, code)
    rows = extract_summary(skeleton_doc(path))
    if not rows:
        print(f"  [건너뜀] {dept}: 제출 업무 없음"); return
    d = Document(path)
    if any("제출 업무 총괄" in p.text for p in d.paragraphs):
        print(f"  [이미 반영됨] {dept}"); return
    before = count_articles(d)
    anchor = next((p for p in d.paragraphs if re.match(r"^제\s*1\s*장", p.text.strip())), None)
    if anchor is None:
        print(f"  [경고] {dept}: 제1장 앵커 없음, 건너뜀"); return

    h = anchor.insert_paragraph_before("")
    rh = h.add_run(f"■ {dept} 제출 업무 총괄  Duties Overview")
    rh.bold = True; rh.font.size = Pt(13); rh.font.name = "맑은 고딕"
    rh.font.color.rgb = RGBColor(0x1A,0x3A,0x5C)
    intro = anchor.insert_paragraph_before("")
    ri = intro.add_run(f"아래 표는 {dept}가 제출한 업무 현황을 반영한 것이다. 세부 운영 지침은 제1장 이하 각 조에서 정한다.")
    ri.font.size = Pt(11); ri.font.name = "맑은 고딕"

    tbl = d.add_table(rows=1+len(rows), cols=3)
    add_borders(tbl)
    for i, htext in enumerate(["업무명","주기","예산"]):
        set_cell(tbl.rows[0].cells[i], htext, bold=True, bg=NAVY, white=True)
    for ri_, (name, cyc, bud) in enumerate(rows, start=1):
        set_cell(tbl.rows[ri_].cells[0], name)
        set_cell(tbl.rows[ri_].cells[1], cyc)
        set_cell(tbl.rows[ri_].cells[2], bud if bud else "—")
    anchor._p.addprevious(tbl._tbl)
    anchor.insert_paragraph_before("")

    d.save(path)
    after = count_articles(Document(path))
    ok = "OK" if after == before else "★조항수변동★"
    print(f"  {dept:8} 업무 {len(rows)}건 주입 | 조항 {before}→{after} {ok}")

if __name__ == "__main__":
    print("제출 업무 총괄 요약표 주입 시작")
    for dept, code in DEPTS:
        inject(dept, code)
    print("완료")
